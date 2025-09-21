import os
import sys
from flask import Flask, request, jsonify, render_template, send_file, session
from dotenv import load_dotenv
from datetime import datetime
from models import db, Project, TeamMember, MeetingSummary, TaskAssignment, MeetingTemplate, User, Role, UserRole
from auth import auth_manager, token_required, admin_required, role_required
# Import modules with error handling
try:
    from validators import InputValidator, MeetingSchema, ProjectSchema, TaskSchema, EmailSchema, UserRegistrationSchema, UserLoginSchema
except ImportError as e:
    print(f"Warning: Could not import validators: {e}")
    InputValidator = None
    MeetingSchema = None
    ProjectSchema = None
    TaskSchema = None
    EmailSchema = None
    UserRegistrationSchema = None
    UserLoginSchema = None

try:
    from rate_limiter import rate_limit, api_rate_limit
except ImportError as e:
    print(f"Warning: Could not import rate_limiter: {e}")
    def rate_limit(requests=100, window=3600):
        def decorator(f):
            return f
        return decorator
    def api_rate_limit():
        def decorator(f):
            return f
        return decorator

try:
    from error_handlers import ErrorHandler, ValidationError, AuthenticationError, AuthorizationError, RateLimitError
except ImportError as e:
    print(f"Warning: Could not import error_handlers: {e}")
    class ErrorHandler:
        def __init__(self, app=None):
            self.app = app
        def init_app(self, app):
            pass
    class ValidationError(Exception):
        pass
    class AuthenticationError(Exception):
        pass
    class AuthorizationError(Exception):
        pass
    class RateLimitError(Exception):
        pass

try:
    from monitoring import init_monitoring
except ImportError as e:
    print(f"Warning: Could not import monitoring: {e}")
    def init_monitoring(app):
        return None

try:
    from database_config import DatabaseConfig
except ImportError as e:
    print(f"Warning: Could not import database_config: {e}")
    class DatabaseConfig:
        def __init__(self, app):
            pass

try:
    from api_docs import register_api_docs
except ImportError as e:
    print(f"Warning: Could not import api_docs: {e}")
    def register_api_docs(app):
        pass

try:
    from notifications import init_notifications
except ImportError as e:
    print(f"Warning: Could not import notifications: {e}")
    def init_notifications(app):
        return None

try:
    from analytics import analytics_engine
except ImportError as e:
    print(f"Warning: Could not import analytics: {e}")
    analytics_engine = None

try:
    from search_engine import search_engine
except ImportError as e:
    print(f"Warning: Could not import search_engine: {e}")
    search_engine = None

try:
    from file_manager import init_file_manager
except ImportError as e:
    print(f"Warning: Could not import file_manager: {e}")
    def init_file_manager(app):
        return None

try:
    from workflow_automation import init_workflow_engine
except ImportError as e:
    print(f"Warning: Could not import workflow_automation: {e}")
    def init_workflow_engine():
        return None
from flask_mail import Mail, Message
import logging
from reportlab.lib.pagesizes import letter, A4
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER, TA_LEFT, TA_JUSTIFY
import io
from knowledge_base import get_app_features
import google.generativeai as genai
from file_processor import process_file
from openpyxl import Workbook
from openpyxl.styles import Font, PatternFill, Alignment
import docx
from docx.shared import Inches
from fpdf import FPDF

load_dotenv()

app = Flask(__name__)

# Security configuration
app.config['SECRET_KEY'] = os.environ.get('SECRET_KEY', 'dev-secret-key-change-in-production')
app.config['JWT_SECRET_KEY'] = os.environ.get('JWT_SECRET_KEY', 'jwt-secret-key-change-in-production')

# Database configuration with proper path handling
db_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'instance', 'meetings.db')
os.makedirs(os.path.dirname(db_path), exist_ok=True)
app.config['SQLALCHEMY_DATABASE_URI'] = f'sqlite:///{db_path}'
app.config['SQLALCHEMY_TRACK_MODIFICATIONS'] = False

# Initialize all components with error handling
try:
    auth_manager.init_app(app)
except Exception as e:
    print(f"Warning: Could not initialize auth_manager: {e}")

try:
    error_handler = ErrorHandler(app)
except Exception as e:
    print(f"Warning: Could not initialize error_handler: {e}")
    error_handler = None

try:
    monitor = init_monitoring(app)
except Exception as e:
    print(f"Warning: Could not initialize monitoring: {e}")
    monitor = None

try:
    db_config = DatabaseConfig(app)
except Exception as e:
    print(f"Warning: Could not initialize database_config: {e}")
    db_config = None

try:
    register_api_docs(app)
except Exception as e:
    print(f"Warning: Could not register API docs: {e}")

try:
    notification_manager = init_notifications(app)
except Exception as e:
    print(f"Warning: Could not initialize notifications: {e}")
    notification_manager = None

try:
    file_manager = init_file_manager(app)
except Exception as e:
    print(f"Warning: Could not initialize file_manager: {e}")
    file_manager = None

try:
    workflow_engine = init_workflow_engine()
except Exception as e:
    print(f"Warning: Could not initialize workflow_engine: {e}")
    workflow_engine = None

# Email configuration with validation
app.config['MAIL_SERVER'] = 'smtp.gmail.com'
app.config['MAIL_PORT'] = 465
app.config['MAIL_USERNAME'] = os.environ.get('MAIL_USERNAME')
app.config['MAIL_PASSWORD'] = os.environ.get('MAIL_PASSWORD')
app.config['MAIL_USE_TLS'] = False
app.config['MAIL_USE_SSL'] = True

db.init_app(app)
mail = Mail(app)

# Configure Google Gemini AI with improved error handling
def initialize_gemini():
    """Initialize Gemini AI with proper error handling"""
    try:
        # Get API key from environment variable
        api_key = os.getenv('GEMINI_API_KEY')
        
        if not api_key:
            logging.warning("GEMINI_API_KEY not found in environment variables")
            # Using provided API key for testing
            api_key = "AIzaSyBS6pnTbIsen01AK1nbqGp-wE8GQfc1rnA"
            if not api_key:
                return None
            
        if len(api_key) < 20:  # Basic validation
            logging.warning("GEMINI_API_KEY appears to be invalid (too short)")
            return None
            
        genai.configure(api_key=api_key)
        model = genai.GenerativeModel('gemini-1.5-flash')
        
        # Test the API key with a simple request
        test_response = model.generate_content("Test")
        logging.info("Gemini AI configured and tested successfully with provided API key")
        return model
        
    except Exception as e:
        logging.error(f"Failed to initialize Gemini AI: {str(e)}")
        return None

# Initialize the model
model = initialize_gemini()

def validate_api_key():
    """Validate and test the Gemini API key configuration"""
    # Using direct API key
    api_key = "AIzaSyDWAKdoiqpzEgDWNu-ir1NRORnXLQ6uMl4"
    
    if not api_key:
        return False, "API key is not set"
    
    if len(api_key) < 20:  # Basic validation
        return False, "API key appears to be invalid (too short)"
    
    if not model:  # Model wasn't initialized due to invalid key
        # Reinitialize model with direct API key
        genai.configure(api_key=api_key)
        model = genai.GenerativeModel('gemini-pro')  # Using gemini-pro for better results
    
    try:
        # Test the API key with a simple request
        test_response = model.generate_content("Test connection")
        return True, "API key is valid and working"
    except Exception as e:
        return False, f"API key test failed: {str(e)}"

@app.route('/api/validate-setup', methods=['GET'])
def validate_setup():
    """Endpoint to validate the application setup"""
    is_valid, message = validate_api_key()
    
    return jsonify({
        "api_key_valid": is_valid,
        "message": message,
        "setup_instructions": {
            "step1": "Visit https://makersuite.google.com/app/apikey",
            "step2": "Create a new API key",
            "step3": "Copy the key and replace 'your_gemini_api_key_here' in the .env file",
            "step4": "Restart the application"
        } if not is_valid else None
    })

logging.basicConfig(level=logging.INFO)


@app.route('/')
def index():
    """Main meeting summarizer page"""
    return render_template('index.html')

# Authentication endpoints
@app.route('/api/auth/register', methods=['POST'])
@rate_limit(requests=5, window=3600)  # 5 registrations per hour
def register():
    """Register a new user"""
    try:
        schema = UserRegistrationSchema()
        data = schema.load(request.get_json())
        
        # Sanitize inputs
        username = InputValidator.sanitize_string(data['username'], 50)
        email = InputValidator.sanitize_string(data['email'], 120)
        
        # Validate email format
        if not InputValidator.validate_email(email):
            return jsonify({'error': 'Invalid email format'}), 400
        
        # Validate password strength
        password_validation = InputValidator.validate_password(data['password'])
        if not password_validation['valid']:
            return jsonify({'error': 'Password validation failed', 'details': password_validation['errors']}), 400
        
        # Create user
        user, message = auth_manager.create_user(
            username=username,
            email=email,
            password=data['password'],
            role=data.get('role', 'user')
        )
        
        if user:
            return jsonify({
                'message': message,
                'user': user.to_dict()
            }), 201
        else:
            return jsonify({'error': message}), 400
            
    except ValidationError as e:
        return jsonify({'error': 'Validation failed', 'details': e.messages}), 400
    except Exception as e:
        logging.error(f"Registration error: {e}")
        return jsonify({'error': 'Registration failed'}), 500

@app.route('/api/auth/login', methods=['POST'])
@rate_limit(requests=10, window=3600)  # 10 login attempts per hour
def login():
    """Authenticate user and return JWT token"""
    try:
        schema = UserLoginSchema()
        data = schema.load(request.get_json())
        
        # Sanitize inputs
        username = InputValidator.sanitize_string(data['username'], 50)
        
        # Authenticate user
        user = auth_manager.authenticate_user(username, data['password'])
        
        if user:
            # Update last login
            user.last_login = datetime.utcnow()
            db.session.commit()
            
            # Generate token
            token = auth_manager.generate_token(user.id, user.username, user.role)
            
            return jsonify({
                'message': 'Login successful',
                'token': token,
                'user': user.to_dict()
            }), 200
        else:
            return jsonify({'error': 'Invalid credentials'}), 401
            
    except ValidationError as e:
        return jsonify({'error': 'Validation failed', 'details': e.messages}), 400
    except Exception as e:
        logging.error(f"Login error: {e}")
        return jsonify({'error': 'Login failed'}), 500

@app.route('/api/auth/me', methods=['GET'])
@token_required
def get_current_user(current_user):
    """Get current user information"""
    return jsonify({'user': current_user.to_dict()}), 200

# Health and monitoring endpoints
@app.route('/api/health', methods=['GET'])
def health_check():
    """Application health check endpoint"""
    try:
        if monitor:
            health_data = monitor.get_health_status()
            status_code = 200 if health_data['status'] == 'healthy' else 503
            return jsonify(health_data), status_code
        else:
            # Basic health check if monitoring is not available
            return jsonify({
                'status': 'healthy',
                'message': 'Application is running (basic mode)',
                'timestamp': datetime.utcnow().isoformat(),
                'monitoring': 'disabled'
            }), 200
    except Exception as e:
        return jsonify({
            'status': 'error',
            'message': f'Health check failed: {str(e)}',
            'timestamp': datetime.utcnow().isoformat()
        }), 500

@app.route('/api/metrics', methods=['GET'])
@token_required
def get_metrics(current_user):
    """Get application performance metrics"""
    try:
        if monitor:
            metrics_data = monitor.get_monitoring_dashboard_data()
            return jsonify(metrics_data), 200
        else:
            # Basic metrics if monitoring is not available
            return jsonify({
                'health': {'status': 'healthy', 'message': 'Basic mode'},
                'performance': {'total_requests': 0, 'average_response_time': 0},
                'errors': {'total_errors': 0},
                'monitoring': 'disabled'
            }), 200
    except Exception as e:
        return jsonify({
            'error': 'Failed to retrieve metrics',
            'message': str(e)
        }), 500


@app.route('/api/summarize', methods=['POST'])
@api_rate_limit()
def summarize_meeting():
    try:
        # Get JSON data
        data = request.get_json()
        if not data:
            return jsonify({'error': 'No JSON data provided'}), 400
        
        # Basic validation
        transcript = data.get('transcript', '').strip()
        client_name = data.get('client_name', '').strip()
        project_name = data.get('project_name', '').strip()
        template_id = data.get('template_id')
        
        # Validate required fields
        if not transcript:
            return jsonify({'error': 'Transcript is required'}), 400
        if not client_name:
            return jsonify({'error': 'Client name is required'}), 400
        if not project_name:
            return jsonify({'error': 'Project name is required'}), 400
        
        # Sanitize inputs if InputValidator is available
        if InputValidator:
            transcript = InputValidator.sanitize_string(transcript, 50000)
            client_name = InputValidator.sanitize_string(client_name, 100)
            project_name = InputValidator.sanitize_string(project_name, 100)
        else:
            # Basic sanitization
            transcript = transcript[:50000]
            client_name = client_name[:100]
            project_name = project_name[:100]

        logging.info(f"Processing: client='{client_name}', project='{project_name}', transcript_length={len(transcript)}")

        if not transcript or not client_name or not project_name:
            error_msg = f"Missing required fields - client: {bool(client_name)}, project: {bool(project_name)}, transcript: {bool(transcript)}"
            logging.warning(error_msg)
            return jsonify({"error": "Please fill in all required fields: Client Name, Project Name, and Meeting Transcript"}), 400

        # Initialize model with direct API key if not initialized
        global model
        if not model:
            api_key = "AIzaSyDWAKdoiqpzEgDWNu-ir1NRORnXLQ6uMl4"
            genai.configure(api_key=api_key)
            model = genai.GenerativeModel('gemini-pro')
            logging.info("Initialized Gemini model with direct API key")

        # Get template if provided
        template_context = ""
        if template_id:
            try:
                template = MeetingTemplate.query.get(template_id)
                if template:
                    template_context = f"\n\nMEETING TEMPLATE: {template.name}\nTEMPLATE FOCUS: {template.default_prompt}\n\nPlease apply this template's focus when analyzing the meeting."
                    logging.info(f"Using meeting template: {template.name}")
            except Exception as e:
                logging.warning(f"Could not load template {template_id}: {e}")
        
        # Use the enhanced AI engine for processing
        try:
            from ai_engine import AIEngine
            ai_engine = AIEngine()
            result = ai_engine.process_transcript(transcript, client_name, project_name)
            logging.info("Successfully processed transcript with enhanced AI engine")
        except Exception as ai_error:
            logging.error(f"Enhanced AI processing failed, falling back to basic processing: {ai_error}")
            # Fallback to basic processing if enhanced fails
            result = create_fallback_analysis(transcript, client_name, project_name)
            logging.info("Using fallback analysis - AI service not available")

        # Find or create project
        try:
            project = Project.query.filter_by(name=project_name, client=client_name).first()
            if not project:
                project = Project(name=project_name, client=client_name)
                db.session.add(project)
                db.session.flush()

            # Save meeting summary
            summary = MeetingSummary(
                transcript=transcript,
                ai_result=result,
                meeting_type="Business Meeting",
                project_id=project.id
            )
            db.session.add(summary)
            
            # Create task assignments from AI result - MVP focus
            if result.get('action_items'):
                for task_data in result['action_items']:
                    # Priority classification
                    original_priority = task_data.get('priority', 'Low').lower()
                    classified_priority = classify_task_priority(
                        task_data.get('task', ''), 
                        transcript  # Use full transcript as context
                    )
                    
                    # Use the more specific classification
                    final_priority = classified_priority if classified_priority != 'normal' else original_priority
                    
                    task = TaskAssignment(
                        task_description=task_data.get('task', ''),
                        priority=final_priority,
                        status='pending',
                        project_id=project.id,
                        meeting_id=summary.id
                    )
                    db.session.add(task)
                    
                    # Send email for high priority tasks
                    if final_priority == 'high':
                        send_critical_task_email(task_data, client_name, project_name)
            
            db.session.commit()
            logging.info(f"Meeting analysis completed for {client_name} - {project_name}")
            
        except Exception as db_error:
            db.session.rollback()
            logging.error(f"Database error during meeting save: {db_error}")
            # Continue and return AI result even if database save fails
        return jsonify(result)
        
    except Exception as e:
        logging.error(f"Error in summarize_meeting: {str(e)}", exc_info=True)
        db.session.rollback()
        return jsonify({"error": f"An internal server error occurred: {str(e)}"}), 500

def create_fallback_analysis(transcript, client_name, project_name):
    """Create a basic analysis when AI is not available"""
    import re
    
    # Extract basic information
    words = transcript.lower().split()
    
    # Simple sentiment analysis
    positive_words = ['good', 'great', 'excellent', 'positive', 'success', 'complete', 'finished']
    negative_words = ['problem', 'issue', 'error', 'failed', 'critical', 'urgent', 'delay']
    
    positive_count = sum(1 for word in words if word in positive_words)
    negative_count = sum(1 for word in words if word in negative_words)
    
    if positive_count > negative_count:
        mood = "Positive"
        mood_justification = "Meeting had more positive language and successful outcomes"
    elif negative_count > positive_count:
        mood = "Negative" 
        mood_justification = "Meeting contained several issues and problems to address"
    else:
        mood = "Neutral"
        mood_justification = "Meeting had balanced discussion of topics"
    
    # Extract action items using simple patterns
    action_items = []
    lines = transcript.split('\n')
    
    for line in lines:
        line_lower = line.lower()
        if any(keyword in line_lower for keyword in ['will', 'need to', 'should', 'must', 'task', 'action']):
            # Extract names (simple pattern)
            names = re.findall(r'[A-Z][a-z]+ [A-Z][a-z]+', line)
            assignee = names[0] if names else 'Unassigned'
            
            # Determine priority
            if any(word in line_lower for word in ['critical', 'urgent', 'asap', 'immediately']):
                priority = 'high'
            elif any(word in line_lower for word in ['important', 'priority', 'soon']):
                priority = 'medium'
            else:
                priority = 'low'
            
            action_items.append({
                'task': line.strip(),
                'assignee': assignee,
                'assigned_by': 'Meeting',
                'deadline': 'Not specified',
                'priority': priority,
                'confidence': 'Medium'
            })
    
    # Extract action items from transcript
    action_items = []
    participants = set()
    
    for line in lines:
        line = line.strip()
        if not line:
            continue
            
        # Extract speaker name (pattern: "Name: content")
        speaker_match = re.match(r'^([A-Z][a-z]+(?:\s+[A-Z][a-z]+)*):\s*(.+)', line)
        if speaker_match:
            speaker = speaker_match.group(1).strip()
            content = speaker_match.group(2).strip()
            participants.add(speaker)
            
            # Look for remark patterns (feedback, comments, suggestions, etc.)
            remark_keywords = [
                'feedback', 'comment', 'suggestion', 'note', 'remark', 'observation',
                'think', 'believe', 'feel', 'consider', 'recommend', 'suggest',
                'concern', 'worry', 'issue', 'problem', 'good', 'great', 'excellent',
                'bad', 'wrong', 'improve', 'better', 'change', 'update'
            ]
            
            content_lower = content.lower()
            if any(keyword in content_lower for keyword in remark_keywords):
                # Try to identify who the remark is directed to
                given_to = 'General'
                for participant in participants:
                    if participant != speaker and participant.lower() in content_lower:
                        given_to = participant
                        break
                
                # Also check for direct addressing patterns
                direct_patterns = [
                    r'to\s+([A-Z][a-z]+)',
                    r'@([A-Z][a-z]+)',
                    r'for\s+([A-Z][a-z]+)',
                    r'([A-Z][a-z]+),?\s+you'
                ]
                
                for pattern in direct_patterns:
                    match = re.search(pattern, content, re.IGNORECASE)
                    if match and match.group(1):
                        given_to = match.group(1)
                        break
                
                action_items.append({
                    'person': speaker,
                    'remark': content,
                    'given_to': given_to
                })
    
    # If no action items found, add a generic one
    if not action_items:
        action_items = [{'person': 'System', 'remark': 'Meeting analysis completed using fallback processing', 'given_to': 'General'}]
    
    # Create summary
    summary = f"Meeting for {client_name} regarding {project_name}. Key topics discussed include project progress, technical implementation, and upcoming deliverables. {len(action_items)} action items identified."
    
    return {
        'summary': summary,
        'mood': {
            'overall': mood,
            'justification': mood_justification
        },
        'action_items': action_items[:5],  # Limit to 5 items
        'key_decisions': [
            'Project timeline and deliverables discussed',
            'Resource allocation planned',
            'Next meeting scheduled'
        ],
        'next_steps': [
            'Follow up on action items',
            'Prepare progress report',
            'Schedule next review meeting'
        ],
        'action_items': action_items[:10]  # Limit to 10 action items
    }

def classify_task_priority(task_text, context=""):
    """Classify task priority based on MVP requirements:
    High: urgent language like 'must', 'ASAP', 'critical'
    Medium: important but not urgent like 'should soon', 'priority'
    Low: standard tasks like 'needs to be done'
    """
    task_lower = (task_text + " " + context).lower()
    
    # High priority indicators
    high_keywords = ['must', 'asap', 'critical', 'urgent', 'immediately', 'emergency', 
                    'deadline', 'crucial', 'vital', 'essential', 'required by']
    
    # Medium priority indicators  
    medium_keywords = ['should soon', 'priority', 'important', 'should', 'preferred',
                      'recommend', 'significant', 'moderate', 'needed soon']
    
    if any(keyword in task_lower for keyword in high_keywords):
        return 'high'
    elif any(keyword in task_lower for keyword in medium_keywords):
        return 'medium'
    else:
        return 'low'

def send_critical_task_email(task_data, client_name, project_name):
    """Send email notification for high priority tasks"""
    try:
        recipients = os.environ.get('NOTIFICATION_RECIPIENTS', '').split(',')
        recipients = [email.strip() for email in recipients if email.strip()]
        
        if not recipients:
            logging.warning("No notification recipients configured")
            return
            
        subject = f"🚨 HIGH PRIORITY TASK ALERT - {client_name} | {project_name}"
        body = f"""
HIGH PRIORITY TASK DETECTED

Client: {client_name}
Project: {project_name}
Task: {task_data.get('task', 'No description')}
Assigned to: {task_data.get('assignee', 'Unassigned')}
Deadline: {task_data.get('deadline', 'No deadline')}

Please review and take action.
        """
        
        msg = Message(
            subject=subject,
            recipients=recipients,
            body=body
        )
        mail.send(msg)
        logging.info(f"High priority task alert sent for: {task_data.get('task', '')[:50]}...")
        
    except Exception as e:
        logging.error(f"Failed to send high priority task email: {e}")

@app.route('/api/projects', methods=['GET', 'POST'])
def handle_projects():
    if request.method == 'GET':
        projects = Project.query.all()
        return jsonify([project.to_dict() for project in projects])
    
    elif request.method == 'POST':
        data = request.get_json()
        project = Project(
            name=data.get('name'),
            client=data.get('client')
        )
        db.session.add(project)
        db.session.commit()
        return jsonify(project.to_dict()), 201

@app.route('/api/team-members', methods=['GET', 'POST'])
def handle_team_members():
    if request.method == 'GET':
        members = TeamMember.query.all()
        return jsonify([member.to_dict() for member in members])
    
    elif request.method == 'POST':
        data = request.get_json()
        member = TeamMember(
            name=data.get('name'),
            role=data.get('role'),
            email=data.get('email')
        )
        db.session.add(member)
        db.session.commit()
        return jsonify(member.to_dict()), 201

@app.route('/api/tasks', methods=['GET', 'POST'])
def handle_tasks():
    if request.method == 'GET':
        tasks = TaskAssignment.query.all()
        return jsonify([task.to_dict() for task in tasks])
    
    elif request.method == 'POST':
        data = request.get_json()
        task = TaskAssignment(
            task_description=data.get('task_description'),
            priority=data.get('priority', 'medium'),
            status=data.get('status', 'pending'),
            project_id=data.get('project_id'),
            assignee_id=data.get('assignee_id')
        )
        db.session.add(task)
        db.session.commit()
        return jsonify(task.to_dict()), 201

@app.route('/api/tasks/<int:task_id>', methods=['PUT'])
def update_task(task_id):
    task = TaskAssignment.query.get_or_404(task_id)
    data = request.get_json()
    
    if 'status' in data:
        task.status = data['status']
    if 'priority' in data:
        task.priority = data['priority']
    if 'assignee_id' in data:
        task.assignee_id = data['assignee_id']
        
    task.updated_at = datetime.utcnow()
    db.session.commit()
    return jsonify(task.to_dict())

@app.route('/api/meetings', methods=['GET'])
def get_meetings():
    meetings = MeetingSummary.query.order_by(MeetingSummary.created_at.desc()).all()
    return jsonify([meeting.to_dict() for meeting in meetings])

@app.route('/api/upload', methods=['POST'])
def upload_file():
    """Handle file uploads and extract text content"""
    try:
        # Check if file is present in request
        if 'file' not in request.files:
            return jsonify({"error": "No file provided"}), 400
        
        file = request.files['file']
        
        # Check if file was selected
        if file.filename == '':
            return jsonify({"error": "No file selected"}), 400
        
        # Validate file extension
        allowed_extensions = ['.pdf', '.docx', '.txt', '.mp3']
        file_extension = '.' + file.filename.split('.')[-1].lower()
        
        if file_extension not in allowed_extensions:
            return jsonify({
                "error": f"Unsupported file type. Allowed: {', '.join(allowed_extensions)}"
            }), 400
        
        # Process the file
        try:
            # Reset file stream position to beginning
            file.stream.seek(0)
            text_content = process_file(file.stream, file.filename)
            
            logging.info(f"Successfully processed file: {file.filename} ({len(text_content)} characters)")
            logging.info(f"Extracted text preview: {text_content[:200]}...")  # Debug log
            
            return jsonify({
                "success": True,
                "transcript": text_content,
                "filename": file.filename,
                "message": f"Successfully processed {file.filename}"
            })
            
        except Exception as processing_error:
            logging.error(f"File processing error for {file.filename}: {processing_error}")
            return jsonify({
                "error": f"Failed to process file: {str(processing_error)}"
            }), 500
            
    except Exception as e:
        logging.error(f"Upload endpoint error: {e}")
        return jsonify({"error": "An internal server error occurred during file upload."}), 500

@app.route('/api/send-mail', methods=['POST'])
def send_mail():
    """Send email to specified recipient with custom message"""
    try:
        data = request.get_json()
        recipient_email = data.get('recipient_email', '').strip()
        message_content = data.get('message', '').strip()
        subject = data.get('subject', 'Message from AI Meeting Summarizer')
        
        if not recipient_email or not message_content:
            return jsonify({"error": "Recipient email and message are required"}), 400
        
        # Validate email format (basic validation)
        if '@' not in recipient_email or '.' not in recipient_email:
            return jsonify({"error": "Invalid email format"}), 400
        
        # Check if email is properly configured
        mail_username = os.environ.get('MAIL_USERNAME')
        mail_password = os.environ.get('MAIL_PASSWORD')
        
        if not mail_username or not mail_password or mail_username == 'your_email@gmail.com' or mail_password == 'your_app_password':
            return jsonify({
                "error": "Email not configured. Please set up Gmail App Password.",
                "setup_required": True,
                "instructions": {
                    "step1": "Go to https://myaccount.google.com/apppasswords",
                    "step2": "Generate App Password for 'Mail'",
                    "step3": "Update MAIL_PASSWORD in .env file",
                    "step4": "Restart the application"
                }
            }), 503
        
        # Create and send email
        msg = Message(
            subject=subject,
            recipients=[recipient_email],
            body=message_content,
            sender=mail_username
        )
        
        mail.send(msg)
        logging.info(f"Email sent successfully to {recipient_email}")
        
        return jsonify({
            "success": True,
            "message": f"Email sent successfully to {recipient_email}"
        })
        
    except Exception as e:
        logging.error(f"Error sending email: {e}")
        error_msg = str(e).lower()
        
        if 'authentication' in error_msg or '535' in error_msg:
            return jsonify({
                "error": "Email authentication failed. Please set up Gmail App Password.",
                "setup_required": True,
                "instructions": {
                    "step1": "Go to https://myaccount.google.com/apppasswords",
                    "step2": "Enable 2-Factor Authentication if not enabled",
                    "step3": "Generate App Password for 'Mail'",
                    "step4": "Replace MAIL_PASSWORD in .env file with 16-character App Password",
                    "step5": "Restart the application"
                }
            }), 503
        else:
            return jsonify({"error": f"Failed to send email: {str(e)}"}), 500


@app.route('/api/assistant', methods=['POST'])
def dictionary_assistant():
    """AI Dictionary Assistant - Provides word definitions and usage examples"""
    try:
        data = request.get_json()
        word = data.get('question', '').strip()
        
        if not word:
            return jsonify({"error": "Please enter a word to look up"}), 400
        
        # Check if AI model is available
        if not model:
            setup_instructions = """🔧 **AI Setup Required**

To use the dictionary, please configure your Google Gemini API key:

1. Create a file named `.env` in the project root directory
2. Copy the contents from `env_template.txt` to `.env`
3. Visit [Google AI Studio](https://makersuite.google.com/app/apikey) to get your API key
4. Replace 'your_gemini_api_key_here' with your actual API key
5. Save the file and restart the application

Example .env file:
```
GEMINI_API_KEY=your_actual_api_key_here
FLASK_ENV=development
FLASK_DEBUG=True
```

Need help? Check the README.md file for detailed setup instructions."""
            
            return jsonify({
                "answer": setup_instructions,
                "mode": "setup_required",
                "success": True
            })
        
        try:
            # Enhanced prompt for comprehensive dictionary definitions
        
            dictionary_prompt = f"""You are an expert dictionary assistant and language consultant. For the word "{word}", provide a comprehensive analysis:

**WORD ANALYSIS FOR: {word.upper()}**

1. **PRONUNCIATION**:
   - Phonetic spelling (IPA format if possible)
   - Syllable breakdown
   - Stress pattern
   - Audio pronunciation guide

2. **DEFINITION**:
   - Primary meaning (clear, concise)
   - Secondary meanings (if applicable)
   - Etymology and word origin
   - Language of origin

3. **GRAMMATICAL INFORMATION**:
   - Part of speech (noun, verb, adjective, etc.)
   - Word forms (plural, past tense, etc.)
   - Grammatical usage notes

4. **PROFESSIONAL USAGE EXAMPLES**:
   - Business/Corporate context
   - Academic/Research context
   - Technical/Professional context
   - Formal communication examples
   - Email/Report writing examples

5. **LANGUAGE FEATURES**:
   - Formality level (formal, informal, neutral)
   - Register (academic, business, casual)
   - Regional variations (if applicable)
   - Common collocations

6. **SYNONYMS & ANTONYMS**:
   - Professional synonyms
   - Academic alternatives
   - Antonyms with context
   - Nuanced differences

7. **STYLISTIC NOTES**:
   - When to use this word
   - When to avoid it
   - Tone and connotation
   - Professional appropriateness

8. **COMMON PHRASES & IDIOMS**:
   - Professional expressions
   - Business terminology
   - Common collocations

Format your response in a clear, structured way that helps professionals understand and use the word effectively in business, academic, and formal contexts. If the word is misspelled, suggest corrections and provide the correct spelling.

Make the response educational, comprehensive, and immediately useful for professional communication."""

            response = model.generate_content(dictionary_prompt)
            answer = response.text.strip()
            
            logging.info(f"Dictionary lookup processed: {word}")
            
            return jsonify({
                "answer": answer,
                "mode": "dictionary",
                "success": True
            })
            
        except Exception as ai_error:
            logging.error(f"AI processing error in dictionary mode: {ai_error}")
            return jsonify({
                "error": "I'm having trouble looking up that word right now. Please try again or check your spelling."
            }), 500
            
    except Exception as e:
        logging.error(f"Error in dictionary assistant: {e}")
        return jsonify({"error": "An internal server error occurred."}), 500

@app.route('/api/export-pdf', methods=['POST'])
def export_task_board_pdf():
    """Export task board to PDF format following memory specification"""
    try:
        data = request.get_json()
        client_name = data.get('client_name', 'Unknown Client')
        project_name = data.get('project_name', 'Unknown Project')
        tasks_by_status = data.get('tasks', {})
        
        # Create PDF
        pdf = FPDF()
        pdf.add_page()
        pdf.set_font('Arial', 'B', 16)
        
        # Header
        pdf.cell(0, 10, 'Task Board Export Report', 0, 1, 'C')
        pdf.ln(5)
        
        pdf.set_font('Arial', '', 12)
        pdf.cell(0, 8, f'Client: {client_name}', 0, 1)
        pdf.cell(0, 8, f'Project: {project_name}', 0, 1)
        pdf.cell(0, 8, f'Export Date: {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}', 0, 1)
        pdf.ln(5)
        
        # Task sections
        sections = [
            ('Pending Tasks', tasks_by_status.get('pending', [])),
            ('In Progress Tasks', tasks_by_status.get('in_progress', [])),
            ('Completed Tasks', tasks_by_status.get('completed', []))
        ]
        
        for section_name, tasks in sections:
            pdf.set_font('Arial', 'B', 14)
            pdf.cell(0, 10, f'{section_name} ({len(tasks)})', 0, 1)
            pdf.set_font('Arial', '', 10)
            
            if not tasks:
                pdf.cell(0, 6, '  No tasks in this category', 0, 1)
            else:
                for i, task in enumerate(tasks, 1):
                    task_text = task.get('task', 'Unnamed Task')
                    priority = task.get('priority', 'normal')
                    assigned_to = task.get('assigned_to', 'Unassigned')
                    
                    pdf.cell(0, 6, f'  {i}. {task_text}', 0, 1)
                    pdf.cell(0, 5, f'     Priority: {priority.title()} | Assigned: {assigned_to}', 0, 1)
            
            pdf.ln(3)
        
        # Save to bytes
        pdf_output = io.BytesIO()
        pdf_content = pdf.output(dest='S').encode('latin1')
        pdf_output.write(pdf_content)
        pdf_output.seek(0)
        
        filename = f'TaskBoard_{client_name.replace(" ", "_")}_{datetime.now().strftime("%Y%m%d_%H%M%S")}.pdf'
        
        return send_file(
            pdf_output,
            mimetype='application/pdf',
            as_attachment=True,
            download_name=filename
        )
        
    except Exception as e:
        logging.error(f"Error generating PDF export: {e}")
        return jsonify({"error": "Failed to generate PDF export"}), 500

@app.route('/api/meeting-templates', methods=['GET', 'POST'])
def handle_meeting_templates():
    """Handle meeting templates CRUD operations"""
    if request.method == 'GET':
        templates = MeetingTemplate.query.filter_by(is_active=True).all()
        return jsonify([template.to_dict() for template in templates])
    
    elif request.method == 'POST':
        data = request.get_json()
        template = MeetingTemplate(
            name=data.get('name'),
            template_type=data.get('template_type'),
            description=data.get('description'),
            default_prompt=data.get('default_prompt')
        )
        db.session.add(template)
        db.session.commit()
        return jsonify(template.to_dict()), 201

@app.route('/api/meeting-templates/<int:template_id>', methods=['GET', 'PUT', 'DELETE'])
def handle_meeting_template(template_id):
    """Handle individual meeting template operations"""
    template = MeetingTemplate.query.get_or_404(template_id)
    
    if request.method == 'GET':
        return jsonify(template.to_dict())
    
    elif request.method == 'PUT':
        data = request.get_json()
        template.name = data.get('name', template.name)
        template.template_type = data.get('template_type', template.template_type)
        template.description = data.get('description', template.description)
        template.default_prompt = data.get('default_prompt', template.default_prompt)
        template.is_active = data.get('is_active', template.is_active)
        db.session.commit()
        return jsonify(template.to_dict())
    
    elif request.method == 'DELETE':
        template.is_active = False
        db.session.commit()
        return jsonify({"message": "Template deactivated successfully"})

@app.route('/api/export-excel', methods=['POST'])
def export_meeting_excel():
    """Export meeting analysis to Excel format"""
    try:
        data = request.get_json()
        client_name = data.get('client_name', 'Unknown Client')
        project_name = data.get('project_name', 'Unknown Project')
        meeting_data = data.get('meeting_data', {})
        
        # Create Excel workbook
        wb = Workbook()
        ws = wb.active
        ws.title = "Meeting Analysis"
        
        # Define styles
        header_font = Font(bold=True, color="FFFFFF")
        header_fill = PatternFill(start_color="366092", end_color="366092", fill_type="solid")
        center_alignment = Alignment(horizontal="center", vertical="center")
        
        # Header information
        ws['A1'] = f"Meeting Analysis Report"
        ws['A1'].font = Font(bold=True, size=16)
        ws.merge_cells('A1:F1')
        
        ws['A3'] = f"Client: {client_name}"
        ws['A4'] = f"Project: {project_name}"
        ws['A5'] = f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}"
        
        # Summary section
        ws['A7'] = "EXECUTIVE SUMMARY"
        ws['A7'].font = header_font
        ws['A7'].fill = header_fill
        ws.merge_cells('A7:F7')
        
        summary = meeting_data.get('summary', 'No summary available')
        ws['A8'] = summary
        ws.merge_cells('A8:F12')
        ws['A8'].alignment = Alignment(wrap_text=True, vertical="top")
        
        # Action Items section
        ws['A14'] = "ACTION ITEMS"
        ws['A14'].font = header_font
        ws['A14'].fill = header_fill
        ws.merge_cells('A14:F14')
        
        # Action items headers
        headers = ['Task', 'Assignee', 'Priority', 'Deadline', 'Status', 'Confidence']
        for col, header in enumerate(headers, 1):
            cell = ws.cell(row=15, column=col, value=header)
            cell.font = Font(bold=True)
            cell.fill = PatternFill(start_color="D9E1F2", end_color="D9E1F2", fill_type="solid")
        
        # Action items data
        action_items = meeting_data.get('action_items', [])
        for row, item in enumerate(action_items, 16):
            if isinstance(item, dict):
                ws.cell(row=row, column=1, value=item.get('task', ''))
                ws.cell(row=row, column=2, value=item.get('assignee', 'Unassigned'))
                ws.cell(row=row, column=3, value=item.get('priority', 'Normal'))
                ws.cell(row=row, column=4, value=item.get('deadline', 'Not specified'))
                ws.cell(row=row, column=5, value='Pending')
                ws.cell(row=row, column=6, value=item.get('confidence', 'Medium'))
            else:
                ws.cell(row=row, column=1, value=str(item))
        
        # Adjust column widths
        ws.column_dimensions['A'].width = 50
        ws.column_dimensions['B'].width = 20
        ws.column_dimensions['C'].width = 15
        ws.column_dimensions['D'].width = 20
        ws.column_dimensions['E'].width = 15
        ws.column_dimensions['F'].width = 15
        
        # Save to bytes
        excel_output = io.BytesIO()
        wb.save(excel_output)
        excel_output.seek(0)
        
        filename = f'MeetingAnalysis_{client_name.replace(" ", "_")}_{datetime.now().strftime("%Y%m%d_%H%M%S")}.xlsx'
        
        return send_file(
            excel_output,
            mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
            as_attachment=True,
            download_name=filename
        )
        
    except Exception as e:
        logging.error(f"Error generating Excel export: {e}")
        return jsonify({"error": "Failed to generate Excel export"}), 500

@app.route('/api/export-word', methods=['POST'])
def export_meeting_word():
    """Export meeting analysis to Word format"""
    try:
        data = request.get_json()
        client_name = data.get('client_name', 'Unknown Client')
        project_name = data.get('project_name', 'Unknown Project')
        meeting_data = data.get('meeting_data', {})
        
        # Create Word document
        doc = docx.Document()
        
        # Title
        title = doc.add_heading('Meeting Analysis Report', 0)
        title.alignment = 1  # Center alignment
        
        # Header information
        doc.add_heading('Meeting Information', level=1)
        doc.add_paragraph(f"Client: {client_name}")
        doc.add_paragraph(f"Project: {project_name}")
        doc.add_paragraph(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
        
        # Executive Summary
        doc.add_heading('Executive Summary', level=1)
        summary = meeting_data.get('summary', 'No summary available')
        doc.add_paragraph(summary)
        
        # Action Items
        doc.add_heading('Action Items', level=1)
        action_items = meeting_data.get('action_items', [])
        
        if action_items:
            # Create table for action items
            table = doc.add_table(rows=1, cols=5)
            table.style = 'Table Grid'
            
            # Header row
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = 'Task'
            hdr_cells[1].text = 'Assignee'
            hdr_cells[2].text = 'Priority'
            hdr_cells[3].text = 'Deadline'
            hdr_cells[4].text = 'Confidence'
            
            # Data rows
            for item in action_items:
                row_cells = table.add_row().cells
                if isinstance(item, dict):
                    row_cells[0].text = item.get('task', '')
                    row_cells[1].text = item.get('assignee', 'Unassigned')
                    row_cells[2].text = item.get('priority', 'Normal')
                    row_cells[3].text = item.get('deadline', 'Not specified')
                    row_cells[4].text = item.get('confidence', 'Medium')
                else:
                    row_cells[0].text = str(item)
        else:
            doc.add_paragraph('No action items identified.')
        
        # Decisions
        decisions = meeting_data.get('key_decisions', [])
        if decisions:
            doc.add_heading('Key Decisions', level=1)
            for decision in decisions:
                doc.add_paragraph(f"• {decision}", style='List Bullet')
        
        # Next Steps
        next_steps = meeting_data.get('next_steps', [])
        if next_steps:
            doc.add_heading('Next Steps', level=1)
            for step in next_steps:
                doc.add_paragraph(f"• {step}", style='List Bullet')
        
        # Save to bytes
        word_output = io.BytesIO()
        doc.save(word_output)
        word_output.seek(0)
        
        filename = f'MeetingAnalysis_{client_name.replace(" ", "_")}_{datetime.now().strftime("%Y%m%d_%H%M%S")}.docx'
        
        return send_file(
            word_output,
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            as_attachment=True,
            download_name=filename
        )
        
    except Exception as e:
        logging.error(f"Error generating Word export: {e}")
        return jsonify({"error": "Failed to generate Word export"}), 500

@app.route('/api/download-meeting-pdf', methods=['POST'])
def download_meeting_pdf():
    """Generate comprehensive PDF report with meeting summary and AI analysis"""
    try:
        data = request.get_json()
        logging.info(f"PDF generation request received")
        
        client_name = data.get('client_name', 'Unknown Client')
        project_name = data.get('project_name', 'Unknown Project')
        transcript = data.get('transcript', '')
        meeting_data = data.get('meeting_data', {})
        
        logging.info(f"Processing PDF for: {client_name} - {project_name}")
        
        # Create PDF using ReportLab
        pdf_output = io.BytesIO()
        doc = SimpleDocTemplate(pdf_output, pagesize=letter, rightMargin=72, leftMargin=72, topMargin=72, bottomMargin=18)
        
        # Get styles
        styles = getSampleStyleSheet()
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Heading1'],
            fontSize=24,
            spaceAfter=30,
            alignment=TA_CENTER,
            textColor=colors.darkblue
        )
        
        heading_style = ParagraphStyle(
            'CustomHeading',
            parent=styles['Heading2'],
            fontSize=16,
            spaceAfter=12,
            textColor=colors.darkblue
        )
        
        normal_style = styles['Normal']
        normal_style.fontSize = 11
        normal_style.spaceAfter = 6
        
        # Build content
        story = []
        
        # Title
        story.append(Paragraph("AI Meeting Analysis Report", title_style))
        story.append(Spacer(1, 20))
        
        # Header Information
        story.append(Paragraph("Meeting Information", heading_style))
        story.append(Paragraph(f"<b>Client:</b> {client_name}", normal_style))
        story.append(Paragraph(f"<b>Project:</b> {project_name}", normal_style))
        story.append(Paragraph(f"<b>Meeting Type:</b> {meeting_data.get('meeting_type', 'Business Meeting')}", normal_style))
        story.append(Paragraph(f"<b>Generated:</b> {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}", normal_style))
        story.append(Spacer(1, 20))
        
        # Participants Section
        if meeting_data.get('participants'):
            story.append(Paragraph("Meeting Participants", heading_style))
            for participant in meeting_data['participants']:
                if isinstance(participant, dict):
                    name = participant.get('name', 'Unknown')
                    role = participant.get('role', 'Not specified')
                    story.append(Paragraph(f"• <b>{name}</b> - {role}", normal_style))
                else:
                    story.append(Paragraph(f"• {participant}", normal_style))
            story.append(Spacer(1, 12))
        
        # Executive Summary
        if meeting_data.get('summary'):
            story.append(Paragraph("Executive Summary", heading_style))
            story.append(Paragraph(str(meeting_data['summary']), normal_style))
            story.append(Spacer(1, 12))
        
        # Mood Analysis
        if meeting_data.get('mood'):
            story.append(Paragraph("Meeting Mood Analysis", heading_style))
            mood = meeting_data['mood']
            story.append(Paragraph(f"<b>Overall Mood:</b> {mood.get('overall', 'Not specified')}", normal_style))
            if mood.get('justification'):
                story.append(Paragraph(f"<b>Analysis:</b> {mood['justification']}", normal_style))
            story.append(Spacer(1, 12))
        
        # Key Decisions
        if meeting_data.get('key_decisions'):
            story.append(Paragraph("Key Decisions Made", heading_style))
            for i, decision in enumerate(meeting_data['key_decisions'], 1):
                story.append(Paragraph(f"{i}. {decision}", normal_style))
            story.append(Spacer(1, 12))
        
        # Action Items
        if meeting_data.get('action_items'):
            story.append(Paragraph("Action Items & Task Assignments", heading_style))
            for i, item in enumerate(meeting_data['action_items'], 1):
                if isinstance(item, dict):
                    story.append(Paragraph(f"<b>{i}. {item.get('task', 'Unnamed Task')}</b>", normal_style))
                    story.append(Paragraph(f"&nbsp;&nbsp;&nbsp;Assigned to: {item.get('assigned_to', 'Unassigned')}", normal_style))
                    story.append(Paragraph(f"&nbsp;&nbsp;&nbsp;Priority: {item.get('priority', 'Normal')}", normal_style))
                    if item.get('deadline'):
                        story.append(Paragraph(f"&nbsp;&nbsp;&nbsp;Deadline: {item['deadline']}", normal_style))
                else:
                    story.append(Paragraph(f"{i}. {item}", normal_style))
            story.append(Spacer(1, 12))
        
        # Next Steps
        if meeting_data.get('next_steps'):
            story.append(Paragraph("Recommended Next Steps", heading_style))
            for i, step in enumerate(meeting_data['next_steps'], 1):
                story.append(Paragraph(f"{i}. {step}", normal_style))
            story.append(Spacer(1, 12))
        
        # Important Action Items
        if meeting_data.get('action_items'):
            story.append(Paragraph("Key Action Items & Quotes", heading_style))
            for remark in meeting_data['action_items']:
                if isinstance(remark, dict):
                    person = remark.get('person', 'Unknown')
                    quote = remark.get('remark', '')
                    story.append(Paragraph(f'<i>"{quote}"</i> - <b>{person}</b>', normal_style))
                else:
                    story.append(Paragraph(f"• {remark}", normal_style))
            story.append(Spacer(1, 12))
        
        # Tasks (if available)
        if meeting_data.get('tasks'):
            story.append(Paragraph("Detailed Task Breakdown", heading_style))
            for i, task in enumerate(meeting_data['tasks'], 1):
                if isinstance(task, dict):
                    story.append(Paragraph(f"<b>{i}. {task.get('task', 'Unnamed Task')}</b>", normal_style))
                    story.append(Paragraph(f"&nbsp;&nbsp;&nbsp;Assigned to: {task.get('assigned_to', 'Unassigned')}", normal_style))
                    story.append(Paragraph(f"&nbsp;&nbsp;&nbsp;Assigned by: {task.get('assigned_by', 'Not specified')}", normal_style))
                    story.append(Paragraph(f"&nbsp;&nbsp;&nbsp;Priority: {task.get('priority', 'Normal')}", normal_style))
                    story.append(Paragraph(f"&nbsp;&nbsp;&nbsp;Confidence: {task.get('confidence', 'Medium')}", normal_style))
                    if task.get('deadline'):
                        story.append(Paragraph(f"&nbsp;&nbsp;&nbsp;Deadline: {task['deadline']}", normal_style))
                else:
                    story.append(Paragraph(f"{i}. {task}", normal_style))
            story.append(Spacer(1, 12))
        
        # Analytics Section
        story.append(Spacer(1, 20))
        story.append(Paragraph("Meeting Analytics", heading_style))
        
        # Calculate basic analytics
        total_action_items = len(meeting_data.get('action_items', []))
        total_decisions = len(meeting_data.get('key_decisions', []))
        total_participants = len(meeting_data.get('participants', []))
        
        # Priority breakdown
        high_priority_tasks = sum(1 for item in meeting_data.get('action_items', []) 
                                if isinstance(item, dict) and item.get('priority', '').lower() in ['high', 'critical'])
        medium_priority_tasks = sum(1 for item in meeting_data.get('action_items', []) 
                                  if isinstance(item, dict) and item.get('priority', '').lower() == 'medium')
        low_priority_tasks = sum(1 for item in meeting_data.get('action_items', []) 
                               if isinstance(item, dict) and item.get('priority', '').lower() == 'low')
        
        story.append(Paragraph(f"<b>Total Action Items:</b> {total_action_items}", normal_style))
        story.append(Paragraph(f"<b>Total Decisions Made:</b> {total_decisions}", normal_style))
        story.append(Paragraph(f"<b>Total Participants:</b> {total_participants}", normal_style))
        story.append(Spacer(1, 8))
        
        story.append(Paragraph("<b>Task Priority Breakdown:</b>", normal_style))
        story.append(Paragraph(f"&nbsp;&nbsp;&nbsp;High Priority: {high_priority_tasks}", normal_style))
        story.append(Paragraph(f"&nbsp;&nbsp;&nbsp;Medium Priority: {medium_priority_tasks}", normal_style))
        story.append(Paragraph(f"&nbsp;&nbsp;&nbsp;Low Priority: {low_priority_tasks}", normal_style))
        story.append(Spacer(1, 12))
        
        # Add transcript if not too long
        if transcript and len(transcript) < 3000:
            story.append(Spacer(1, 20))
            story.append(Paragraph("Meeting Transcript", heading_style))
            story.append(Paragraph(str(transcript), normal_style))
        
        # Footer
        story.append(Spacer(1, 30))
        footer_style = ParagraphStyle(
            'Footer',
            parent=styles['Normal'],
            fontSize=8,
            alignment=TA_CENTER,
            textColor=colors.grey
        )
        story.append(Paragraph(f"Generated by AI Meeting Summarizer - {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}", footer_style))
        
        # Build PDF
        doc.build(story)
        pdf_output.seek(0)
        
        filename = f'Meeting_Report_{client_name.replace(" ", "_")}_{project_name.replace(" ", "_")}_{datetime.now().strftime("%Y%m%d_%H%M%S")}.pdf'
        
        logging.info(f"PDF generated successfully, size: {len(pdf_output.getvalue())} bytes")
        
        return send_file(
            pdf_output,
            mimetype='application/pdf',
            as_attachment=True,
            download_name=filename
        )
        
    except Exception as e:
        logging.error(f"Error generating meeting PDF: {e}")
        return jsonify({"error": f"Failed to generate PDF report: {str(e)}"}), 500

@app.route('/api/analytics', methods=['GET'])
def get_meeting_analytics():
    """Get comprehensive meeting analytics and insights"""
    try:
        # Get all meetings
        meetings = MeetingSummary.query.order_by(MeetingSummary.created_at.desc()).all()
        
        analytics = {
            'total_meetings': len(meetings),
            'total_projects': len(set(m.project_id for m in meetings)),
            'total_tasks': TaskAssignment.query.count(),
            'completed_tasks': TaskAssignment.query.filter_by(status='completed').count(),
            'pending_tasks': TaskAssignment.query.filter_by(status='pending').count(),
            'critical_tasks': TaskAssignment.query.filter_by(priority='critical').count(),
            'average_meeting_effectiveness': 0,
            'meeting_types': {},
            'sentiment_distribution': {'Positive': 0, 'Negative': 0, 'Neutral': 0},
            'recent_meetings': []
        }
        
        # Calculate analytics from meeting data
        effectiveness_scores = []
        for meeting in meetings:
            ai_result = meeting.ai_result
            if isinstance(ai_result, dict):
                # Meeting effectiveness
                if 'meeting_effectiveness_score' in ai_result:
                    effectiveness_scores.append(ai_result['meeting_effectiveness_score'])
                
                # Meeting types
                meeting_type = ai_result.get('meeting_type', 'General')
                analytics['meeting_types'][meeting_type] = analytics['meeting_types'].get(meeting_type, 0) + 1
                
                # Sentiment distribution
                sentiment = ai_result.get('sentiment_analysis', {}).get('overall_sentiment', 'Neutral')
                if sentiment in analytics['sentiment_distribution']:
                    analytics['sentiment_distribution'][sentiment] += 1
                
                # Recent meetings
                if len(analytics['recent_meetings']) < 5:
                    analytics['recent_meetings'].append({
                        'id': meeting.id,
                        'project': meeting.project.name if meeting.project else 'Unknown',
                        'client': meeting.project.client if meeting.project else 'Unknown',
                        'created_at': meeting.created_at.isoformat(),
                        'meeting_type': meeting_type,
                        'effectiveness_score': ai_result.get('meeting_effectiveness_score', 0)
                    })
        
        # Calculate average effectiveness
        if effectiveness_scores:
            analytics['average_meeting_effectiveness'] = sum(effectiveness_scores) / len(effectiveness_scores)
        
        return jsonify(analytics)
        
    except Exception as e:
        logging.error(f"Error generating analytics: {e}")
        return jsonify({"error": "Failed to generate analytics"}), 500

def initialize_database():
    """Initialize database with proper error handling"""
    try:
        with app.app_context():
            # Create all database tables
            db.create_all()
            logging.info("Database tables created successfully")
            
            # Create default team members if none exist
            if TeamMember.query.count() == 0:
                default_members = [
                    TeamMember(name="Sarah Johnson", role="Project Manager", email="sarah.johnson@company.com"),
                    TeamMember(name="Mike Chen", role="Senior Developer", email="mike.chen@company.com"),
                    TeamMember(name="Lisa Martinez", role="UX Designer", email="lisa.martinez@company.com"),
                    TeamMember(name="David Brown", role="Business Analyst", email="david.brown@company.com"),
                    TeamMember(name="Emma Wilson", role="QA Engineer", email="emma.wilson@company.com")
                ]
                for member in default_members:
                    db.session.add(member)
                db.session.commit()
                logging.info("Created default team members")
            
            # Create default meeting templates if none exist
            if MeetingTemplate.query.count() == 0:
                default_templates = [
                    MeetingTemplate(
                        name="Sales Call",
                        template_type="sales",
                        description="Template for sales meetings and client calls",
                        default_prompt="Focus on sales objectives, client needs, objections, and next steps. Prioritize deal progression and relationship building."
                    ),
                    MeetingTemplate(
                        name="Project Review",
                        template_type="project",
                        description="Template for project status and milestone reviews",
                        default_prompt="Emphasize project progress, timeline adherence, resource allocation, and risk mitigation. Focus on deliverables and deadlines."
                    ),
                    MeetingTemplate(
                        name="Client Meeting",
                        template_type="client",
                        description="Template for client-facing meetings and presentations",
                        default_prompt="Highlight client satisfaction, requirements gathering, feedback collection, and service delivery. Focus on client value and expectations."
                    ),
                    MeetingTemplate(
                        name="Internal Team",
                        template_type="internal",
                        description="Template for internal team meetings and standups",
                        default_prompt="Focus on team coordination, task assignments, blockers, and collaboration. Emphasize productivity and team dynamics."
                    )
                ]
                for template in default_templates:
                    db.session.add(template)
                db.session.commit()
                logging.info("Created default meeting templates")
    except Exception as e:
        logging.error(f"Database initialization error: {e}")
        sys.exit(1)

def check_dependencies():
    """Check if all required dependencies are available"""
    required_modules = [
        'flask', 'flask_sqlalchemy', 'flask_mail', 
        'google.generativeai', 'dotenv', 'fpdf'
    ]
    
    missing_modules = []
    for module in required_modules:
        try:
            __import__(module)
        except ImportError:
            missing_modules.append(module)
    
    if missing_modules:
        logging.error(f"Missing required modules: {', '.join(missing_modules)}")
        logging.error("Please install missing dependencies with: pip install -r requirements.txt")
        return False
    
    return True

# Advanced Analytics Endpoints
@app.route('/api/analytics/dashboard', methods=['GET'])
@token_required
def get_analytics_dashboard(current_user):
    """Get comprehensive analytics dashboard"""
    try:
        days = request.args.get('days', 30, type=int)
        project_id = request.args.get('project_id', type=int)
        
        metrics = analytics_engine.get_dashboard_metrics(
            user_id=current_user.id,
            project_id=project_id,
            days=days
        )
        
        return jsonify(metrics), 200
        
    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/api/search/global', methods=['GET'])
@token_required
def global_search(current_user):
    """Global search across all entities"""
    try:
        query = request.args.get('q', '')
        limit = request.args.get('limit', 20, type=int)
        
        if not query:
            return jsonify({'error': 'Search query required'}), 400
        
        results = search_engine.global_search(
            query=query,
            user_id=current_user.id,
            limit=limit
        )
        
        return jsonify(results), 200
        
    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/api/files/upload', methods=['POST'])
@token_required
def upload_file_advanced(current_user):
    """Upload file with advanced processing"""
    try:
        if 'file' not in request.files:
            return jsonify({'error': 'No file provided'}), 400
        
        file = request.files['file']
        category = request.form.get('category', 'documents')
        
        result = file_manager.upload_file(
            file=file,
            user_id=current_user.id,
            category=category
        )
        
        return jsonify(result), 200 if result['success'] else 400
        
    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/api/workflows', methods=['GET'])
@token_required
def get_workflows(current_user):
    """Get available workflows"""
    try:
        workflows = []
        for workflow_id, workflow in workflow_engine.workflows.items():
            workflows.append({
                'id': workflow_id,
                'name': workflow.get('name'),
                'description': workflow.get('description'),
                'enabled': workflow.get('enabled', True)
            })
        
        return jsonify({'workflows': workflows}), 200
        
    except Exception as e:
        return jsonify({'error': str(e)}), 500

if __name__ == '__main__':
    # Configure logging
    logging.basicConfig(
        level=logging.INFO,
        format='%(asctime)s - %(name)s - %(levelname)s - %(message)s',
        handlers=[
            logging.FileHandler('app.log'),
            logging.StreamHandler(sys.stdout)
        ]
    )
    
    logging.info("Starting AI Meeting Summarizer...")
    
    # Check dependencies
    if not check_dependencies():
        logging.error("Dependency check failed. Exiting.")
        sys.exit(1)
    
    # Initialize database (simplified for quick startup)
    try:
        with app.app_context():
            db.create_all()
            logging.info("Database tables created successfully")
            
            # Create default admin user if none exists
            if User.query.count() == 0:
                admin_user = User(
                    name="Admin User",
                    email="admin@example.com",
                    role="admin"
                )
                admin_user.set_password("admin123")
                db.session.add(admin_user)
                db.session.commit()
                logging.info("Created default admin user: admin@example.com / admin123")
    except Exception as e:
        logging.error(f"Database initialization error: {e}")
        logging.warning("Continuing without full database initialization...")
    
    # Check AI configuration
    if model:
        logging.info("✅ AI Meeting Summarizer ready with AI capabilities")
    else:
        logging.warning("⚠️ AI Meeting Summarizer running in limited mode (no AI)")
        logging.warning("   To enable AI features, configure your GEMINI_API_KEY")
    
    logging.info("🚀 Starting Flask application on http://localhost:5000")
    logging.info("📈 Analytics: http://localhost:5000/api/analytics/dashboard")
    logging.info("🔍 Search: http://localhost:5000/api/search/global")
    logging.info("📁 Files: http://localhost:5000/api/files/upload")
    logging.info("⚙️ Workflows: http://localhost:5000/api/workflows")
    app.run(debug=True, port=5000, host='0.0.0.0')